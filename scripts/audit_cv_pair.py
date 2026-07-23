#!/usr/bin/env python3
"""Audit English/Finnish CV parity, links, privacy, and PDF basics."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from pypdf import PdfReader


EXPECTED_URLS = {
    "mailto:contact@zyinnova.com",
    "https://portfolio.zyinnova.com/",
    "https://github.com/ZiyaaddinYaramis",
    "https://zyinnova.com/",
    "https://zyinnova.com/projects/lemmikkilife",
    "https://github.com/ZiyaaddinYaramis/github_io_todo_app",
    "https://github.com/ZiyaaddinYaramis/portfolio-zyinnova-portfolio",
}

SHARED_FACTS = [
    "Ziyaaddin Yaramis",
    "ZyInnova",
    "LemmikkiLife",
    "GameScan",
    "Flutter",
    "Dart",
    "Firebase Authentication",
    "Cloud Firestore",
    "REST/HTTP",
    "TMDB API",
    "HTML",
    "CSS",
    "JavaScript",
    "2025",
    "Roseance Oy",
    "Sisustusklinikka.fi",
    "Wise Quarter",
    "TechPro Education",
    "2006-2016",
]

PROJECT_ORDER = ["LemmikkiLife", "Todo", "Portfolio"]

EN_EXPERIENCE_FACTS = [
    "Earlier Experience",
    "Flutter Developer Intern",
    "Oct 2024-Dec 2024",
    "Web Developer Intern",
    "Feb 2024-May 2024",
    "Java Mentor",
    "Police Officer",
    "Turkish National Police",
    "Additional Training",
    "Full-Stack Java Developer Training",
    "Completed 2021",
]

FI_EXPERIENCE_FACTS = [
    "Aiempi työkokemus",
    "Flutter-kehittäjäharjoittelija",
    "loka-joulukuu 2024",
    "Web-kehittäjäharjoittelija",
    "helmi-toukokuu 2024",
    "Java-mentori",
    "Poliisi",
    "Turkin kansallinen poliisi",
    "Lisäkoulutus",
    "Full-Stack Java Developer -koulutus",
    "Suoritettu 2021",
]


def docx_text_and_urls(path: Path) -> tuple[str, set[str], Document]:
    document = Document(path)
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    urls = {
        rel.target_ref
        for rel in document.part.rels.values()
        if rel.reltype == RT.HYPERLINK
    }
    return text, urls, document


def pdf_text_and_urls(path: Path) -> tuple[str, set[str], PdfReader]:
    reader = PdfReader(path)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    urls: set[str] = set()
    for page in reader.pages:
        for annotation_ref in page.get("/Annots", []):
            annotation = annotation_ref.get_object()
            action = annotation.get("/A")
            if action and action.get("/URI"):
                urls.add(str(action["/URI"]))
    return text, urls, reader


def ordered(text: str, values: list[str]) -> bool:
    cursor = 0
    for value in values:
        position = text.find(value, cursor)
        if position < 0:
            return False
        cursor = position + len(value)
    return True


def no_private_contact(text: str) -> bool:
    phone = re.compile(
        r"(?<!\w)(?:\+\d[\d ()-]{7,}\d|\d{2,4}[ ()]\d[\d ()-]{5,}\d)"
    )
    return "@gmail.com" not in text.lower() and not phone.search(text)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--en-docx", type=Path, required=True)
    parser.add_argument("--fi-docx", type=Path, required=True)
    parser.add_argument("--en-pdf", type=Path, required=True)
    parser.add_argument("--fi-pdf", type=Path, required=True)
    parser.add_argument("--out", type=Path)
    args = parser.parse_args()

    en_docx_text, en_docx_urls, en_docx = docx_text_and_urls(args.en_docx)
    fi_docx_text, fi_docx_urls, fi_docx = docx_text_and_urls(args.fi_docx)
    en_pdf_text, en_pdf_urls, en_pdf = pdf_text_and_urls(args.en_pdf)
    fi_pdf_text, fi_pdf_urls, fi_pdf = pdf_text_and_urls(args.fi_pdf)

    checks = {
        "docx_shared_facts_en": all(
            fact in en_docx_text for fact in SHARED_FACTS
        ),
        "docx_shared_facts_fi": all(
            fact in fi_docx_text for fact in SHARED_FACTS
        ),
        "pdf_shared_facts_en": all(
            fact in en_pdf_text for fact in SHARED_FACTS
        ),
        "pdf_shared_facts_fi": all(
            fact in fi_pdf_text for fact in SHARED_FACTS
        ),
        "project_order_en": ordered(en_docx_text, PROJECT_ORDER),
        "project_order_fi": ordered(fi_docx_text, PROJECT_ORDER),
        "language_levels_en": all(
            value in en_docx_text
            for value in ["Turkish", "Native", "English", "Finnish", "B2"]
        ),
        "language_levels_fi": all(
            value in fi_docx_text
            for value in ["Turkki", "äidinkieli", "Englanti", "Suomi", "B2"]
        ),
        "experience_facts_en": all(
            value.casefold() in en_docx_text.casefold()
            and value.casefold() in en_pdf_text.casefold()
            for value in EN_EXPERIENCE_FACTS
        ),
        "experience_facts_fi": all(
            value.casefold() in fi_docx_text.casefold()
            and value.casefold() in fi_pdf_text.casefold()
            for value in FI_EXPERIENCE_FACTS
        ),
        "docx_links_en": en_docx_urls == EXPECTED_URLS,
        "docx_links_fi": fi_docx_urls == EXPECTED_URLS,
        "pdf_links_en": en_pdf_urls == EXPECTED_URLS,
        "pdf_links_fi": fi_pdf_urls == EXPECTED_URLS,
        "no_private_contact_en": no_private_contact(en_docx_text + en_pdf_text),
        "no_private_contact_fi": no_private_contact(fi_docx_text + fi_pdf_text),
        "one_or_two_pages_en": 1 <= len(en_pdf.pages) <= 2,
        "one_or_two_pages_fi": 1 <= len(fi_pdf.pages) <= 2,
        "no_tables_en": not en_docx.tables,
        "no_tables_fi": not fi_docx.tables,
        "blank_docx_author_en": not en_docx.core_properties.author,
        "blank_docx_author_fi": not fi_docx.core_properties.author,
    }
    result = {
        "passed": all(checks.values()),
        "checks": checks,
        "character_counts": {
            "en_pdf": len(en_pdf_text),
            "fi_pdf": len(fi_pdf_text),
        },
        "unique_links": {
            "en_docx": sorted(en_docx_urls),
            "fi_docx": sorted(fi_docx_urls),
            "en_pdf": sorted(en_pdf_urls),
            "fi_pdf": sorted(fi_pdf_urls),
        },
    }
    output = json.dumps(result, ensure_ascii=False, indent=2)
    if args.out:
        args.out.parent.mkdir(parents=True, exist_ok=True)
        args.out.write_text(output + "\n", encoding="utf-8")
    print(output)
    raise SystemExit(0 if result["passed"] else 1)


if __name__ == "__main__":
    main()
