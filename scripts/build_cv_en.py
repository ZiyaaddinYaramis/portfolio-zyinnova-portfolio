#!/usr/bin/env python3
"""Build the general English CV as a deterministic DOCX.

Design authority:
- Base preset: compact_reference_guide.
- Header pattern: customer_pack, adapted to a resume title/contact stack.
- Named override: European CV page size uses A4 portrait.
- Named override: CV density uses 10.2 pt body, 1.10 line spacing, compact
  section rhythm, and 0.62-0.72 inch margins for a concise one-to-two-page CV.
- No tables, images, text boxes, or layout-only shapes are used.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Cm, Inches, Pt, RGBColor


FONT = "Calibri"
BODY_COLOR = RGBColor(31, 47, 59)
NAVY = RGBColor(11, 37, 69)
ACCENT = RGBColor(18, 112, 130)
MUTED = RGBColor(89, 102, 113)
LINK = RGBColor(0, 101, 130)


def set_run_font(
    run,
    *,
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_free_page_geometry(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Inches(0.62)
    section.right_margin = Inches(0.72)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.72)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(10.2)
    normal.font.color.rgb = BODY_COLOR
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    heading = document.styles["Heading 1"]
    heading.font.name = FONT
    heading._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    heading.font.size = Pt(12)
    heading.font.bold = True
    heading.font.color.rgb = ACCENT
    heading.paragraph_format.space_before = Pt(7)
    heading.paragraph_format.space_after = Pt(3)
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.keep_together = True

    entry = document.styles.add_style("CV Entry Title", WD_STYLE_TYPE.PARAGRAPH)
    entry.base_style = normal
    entry.font.name = FONT
    entry._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    entry._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    entry.font.size = Pt(10.5)
    entry.font.bold = True
    entry.font.color.rgb = NAVY
    entry.paragraph_format.space_before = Pt(2)
    entry.paragraph_format.space_after = Pt(0)
    entry.paragraph_format.keep_with_next = True

    meta = document.styles.add_style("CV Metadata", WD_STYLE_TYPE.PARAGRAPH)
    meta.base_style = normal
    meta.font.name = FONT
    meta._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    meta._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    meta.font.size = Pt(9.0)
    meta.font.color.rgb = MUTED
    meta.paragraph_format.space_before = Pt(0)
    meta.paragraph_format.space_after = Pt(2)
    meta.paragraph_format.keep_with_next = True

    compact = document.styles.add_style("CV Compact", WD_STYLE_TYPE.PARAGRAPH)
    compact.base_style = normal
    compact.font.name = FONT
    compact._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    compact._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    compact.font.size = Pt(10.0)
    compact.font.color.rgb = BODY_COLOR
    compact.paragraph_format.space_before = Pt(0)
    compact.paragraph_format.space_after = Pt(1)
    compact.paragraph_format.line_spacing = 1.08


def create_bullet_numbering(document: Document) -> int:
    numbering = document.part.numbering_part.element
    abstract_ids = [
        int(item.get(qn("w:abstractNumId")))
        for item in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [
        int(item.get(qn("w:numId")))
        for item in numbering.findall(qn("w:num"))
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level.append(num_fmt)

    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)

    level_jc = OxmlElement("w:lvlJc")
    level_jc.set(qn("w:val"), "left")
    level.append(level_jc)

    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)

    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(indent)

    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "264")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    level.append(ppr)

    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    rpr.append(fonts)
    level.append(rpr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_bullet(paragraph, num_id: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    paragraph.paragraph_format.widow_control = True


def add_hyperlink(paragraph, text: str, url: str, *, bold: bool = False) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    rpr.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), f"{LINK[0]:02X}{LINK[1]:02X}{LINK[2]:02X}")
    rpr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "19")
    rpr.append(size)
    if bold:
        rpr.append(OxmlElement("w:b"))
    run.append(rpr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_contact_line(document: Document) -> None:
    paragraph = document.add_paragraph(style="CV Metadata")
    add_hyperlink(
        paragraph, "contact@zyinnova.com", "mailto:contact@zyinnova.com"
    )
    for label, url in [
        ("Portfolio", "https://portfolio.zyinnova.com/"),
        ("GitHub", "https://github.com/ZiyaaddinYaramis"),
        ("ZyInnova", "https://zyinnova.com/"),
    ]:
        separator = paragraph.add_run("  |  ")
        set_run_font(separator, size=9.0, color=MUTED)
        add_hyperlink(paragraph, label, url)


def add_section(document: Document, title: str) -> None:
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(title.upper())
    set_run_font(run, size=12, color=ACCENT, bold=True)


def add_entry_title(document: Document, title: str, status: str) -> None:
    paragraph = document.add_paragraph(style="CV Entry Title")
    title_run = paragraph.add_run(title)
    set_run_font(title_run, size=10.5, color=NAVY, bold=True)
    status_run = paragraph.add_run(f"  |  {status}")
    set_run_font(status_run, size=10.0, color=MUTED, bold=False)


def add_bullet(document: Document, text: str, num_id: int) -> None:
    paragraph = document.add_paragraph()
    apply_bullet(paragraph, num_id)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.2, color=BODY_COLOR)


def add_labeled_skill(
    document: Document, label: str, value: str, num_id: int
) -> None:
    paragraph = document.add_paragraph()
    apply_bullet(paragraph, num_id)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=10.2, color=NAVY, bold=True)
    value_run = paragraph.add_run(value)
    set_run_font(value_run, size=10.2, color=BODY_COLOR)


def add_project_link(
    document: Document, label: str, text: str, url: str
) -> None:
    paragraph = document.add_paragraph(style="CV Metadata")
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, size=9.0, color=MUTED, bold=True)
    add_hyperlink(paragraph, text, url)


def add_page_field(paragraph) -> None:
    for tag, value in [
        ("w:fldChar", "begin"),
        ("w:instrText", " PAGE "),
        ("w:fldChar", "end"),
    ]:
        run = OxmlElement("w:r")
        run_props = OxmlElement("w:rPr")
        color = OxmlElement("w:color")
        color.set(qn("w:val"), "596671")
        run_props.append(color)
        size = OxmlElement("w:sz")
        size.set(qn("w:val"), "17")
        run_props.append(size)
        run.append(run_props)
        node = OxmlElement(tag)
        if tag == "w:fldChar":
            node.set(qn("w:fldCharType"), value)
        else:
            node.set(qn("xml:space"), "preserve")
            node.text = value
        run.append(node)
        paragraph._p.append(run)


def configure_footer(document: Document) -> None:
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Ziyaaddin Yaramis | English CV | ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(paragraph)


def build_document() -> Document:
    document = Document()
    set_cell_free_page_geometry(document)
    configure_styles(document)
    configure_footer(document)
    bullet_num_id = create_bullet_numbering(document)

    properties = document.core_properties
    properties.title = "Ziyaaddin Yaramis - English CV"
    properties.subject = "General English curriculum vitae"
    properties.keywords = "Software Developer, ZyInnova, Flutter, Firebase"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.created = datetime(2026, 7, 23, 12, 0, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 23, 12, 0, tzinfo=timezone.utc)
    properties.revision = 1

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    title.paragraph_format.keep_with_next = True
    name = title.add_run("Ziyaaddin Yaramis")
    set_run_font(name, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.paragraph_format.keep_with_next = True
    role = subtitle.add_run(
        "Full-Stack Software Developer | Founder, ZyInnova"
    )
    set_run_font(role, size=12, color=ACCENT, bold=True)
    add_contact_line(document)

    add_section(document, "Professional Summary")
    summary = document.add_paragraph()
    summary_run = summary.add_run(
        "Software developer and founder of ZyInnova, building practical web "
        "and mobile product work with an emphasis on clear scope, maintainable "
        "implementation, and evidence-based communication. Hands-on project "
        "experience includes Flutter, Firebase Authentication, Cloud "
        "Firestore, REST/HTTP integrations, and HTML/CSS/JavaScript. Works "
        "professionally in English and Finnish at B2 level."
    )
    set_run_font(summary_run, size=10.2, color=BODY_COLOR)

    add_section(document, "Core Skills")
    for label, value in [
        ("Mobile", "Flutter, Dart"),
        ("Web", "HTML, CSS, JavaScript, Bootstrap utilities"),
        ("Backend and data", "Firebase Authentication, Cloud Firestore"),
        ("Integration", "REST/HTTP APIs, TMDB API"),
        ("Workflow", "Git, GitHub, documentation, iterative delivery"),
    ]:
        add_labeled_skill(document, label, value, bullet_num_id)

    add_section(document, "Experience")
    add_entry_title(
        document, "Founder & Software Developer - ZyInnova", "Current"
    )
    for text in [
        "Define and develop founder-led web and mobile product work under the "
        "ZyInnova umbrella, from product scope through implementation and "
        "iteration.",
        "Maintain the public founder portfolio and keep project communication "
        "aligned with the current development status; current private product "
        "work includes LemmikkiLife and GameScan.",
    ]:
        add_bullet(document, text, bullet_num_id)
    add_project_link(
        document, "Website", "ZyInnova website", "https://zyinnova.com/"
    )

    add_section(document, "Earlier Experience")
    earlier_experience = [
        (
            "Flutter Developer Intern - Roseance Oy",
            "Oct 2024-Dec 2024",
            "Completed a documented fixed-term internship in a Flutter "
            "developer role.",
        ),
        (
            "Web Developer Intern - Sisustusklinikka.fi",
            "Feb 2024-May 2024",
            "Completed a documented fixed-term internship in a web developer "
            "role.",
        ),
        (
            "Java Mentor - Wise Quarter",
            "2022",
            "Served as a Java mentor in a documented training program.",
        ),
        (
            "Police Officer - Turkish National Police",
            "2006-2016",
            "Served as a police officer in the Turkish National Police.",
        ),
    ]
    for title_text, period, description in earlier_experience:
        add_entry_title(document, title_text, period)
        add_bullet(document, description, bullet_num_id)

    document.add_page_break()
    add_section(document, "Selected Projects")

    add_entry_title(document, "LemmikkiLife", "In development")
    for text in [
        "Developing a private ZyInnova pet-care product intended to bring care "
        "tracking, vaccination and medication reminders, premium products, "
        "and subscription boxes into one application.",
    ]:
        add_bullet(document, text, bullet_num_id)
    add_project_link(
        document,
        "Project page",
        "LemmikkiLife project page",
        "https://zyinnova.com/projects/lemmikkilife",
    )

    add_entry_title(document, "Multi-language Todo App", "Public prototype")
    for text in [
        "Built a Flutter application with English and Finnish localization, "
        "Firebase Authentication, password reset, user-specific Cloud "
        "Firestore task data, priorities, due dates, filtering, and local "
        "notifications.",
    ]:
        add_bullet(document, text, bullet_num_id)
    add_project_link(
        document,
        "Repository",
        "Todo App repository",
        "https://github.com/ZiyaaddinYaramis/github_io_todo_app",
    )

    add_entry_title(document, "Founder Portfolio", "Live")
    for text in [
        "Built and published the canonical personal portfolio using HTML, "
        "CSS, and JavaScript with Bootstrap utilities, AOS, and Typed.js; "
        "maintains the public professional profile at portfolio.zyinnova.com.",
    ]:
        add_bullet(document, text, bullet_num_id)
    add_project_link(
        document,
        "Live site",
        "Founder Portfolio",
        "https://portfolio.zyinnova.com/",
    )
    add_project_link(
        document,
        "Repository",
        "Portfolio repository",
        "https://github.com/ZiyaaddinYaramis/portfolio-zyinnova-portfolio",
    )

    add_section(document, "Education")
    add_entry_title(
        document,
        "Vocational Qualification in Information and Communication Technology",
        "Graduated 2025",
    )
    education = document.add_paragraph(style="CV Compact")
    institution = education.add_run("Keuda, Finland")
    set_run_font(institution, size=10.0, color=BODY_COLOR, bold=True)
    qualification = education.add_run(" | Qualification title: Software Developer")
    set_run_font(qualification, size=10.0, color=BODY_COLOR)

    add_section(document, "Additional Training")
    add_entry_title(
        document,
        "Full-Stack Java Developer Training - TechPro Education",
        "Completed 2021",
    )

    add_section(document, "Languages")
    languages = document.add_paragraph(style="CV Compact")
    for index, (language, level) in enumerate(
        [("Turkish", "Native"), ("English", "B2"), ("Finnish", "B2")]
    ):
        if index:
            separator = languages.add_run("  |  ")
            set_run_font(separator, size=10.2, color=MUTED)
        label = languages.add_run(f"{language}: ")
        set_run_font(label, size=10.2, color=NAVY, bold=True)
        value = languages.add_run(level)
        set_run_font(value, size=10.2, color=BODY_COLOR)

    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(args.output)


if __name__ == "__main__":
    main()
