#!/usr/bin/env python3
"""Build the general Finnish CV as a deterministic DOCX.

The Finnish version uses the same compact single-column A4 design as the
verified English CV. Content is adapted from the same canonical facts.
"""

from __future__ import annotations

import argparse
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from build_cv_en import (
    ACCENT,
    BODY_COLOR,
    MUTED,
    NAVY,
    add_bullet,
    add_contact_line,
    add_entry_title,
    add_labeled_skill,
    add_page_field,
    add_project_link,
    add_section,
    configure_styles,
    create_bullet_numbering,
    set_cell_free_page_geometry,
    set_run_font,
)


def configure_footer(document: Document) -> None:
    paragraph = document.sections[0].footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Ziyaaddin Yaramis | Suomenkielinen CV | ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(paragraph)


def build_document() -> Document:
    document = Document()
    set_cell_free_page_geometry(document)
    configure_styles(document)
    configure_footer(document)
    bullet_num_id = create_bullet_numbering(document)

    properties = document.core_properties
    properties.title = "Ziyaaddin Yaramis - Suomenkielinen CV"
    properties.subject = "Yleinen suomenkielinen ansioluettelo"
    properties.keywords = "Ohjelmistokehittäjä, ZyInnova, Flutter, Firebase"
    properties.author = ""
    properties.last_modified_by = ""
    properties.comments = ""
    properties.created = datetime(2026, 7, 23, 12, 0, tzinfo=timezone.utc)
    properties.modified = datetime(2026, 7, 23, 12, 0, tzinfo=timezone.utc)
    properties.revision = 1

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    title.paragraph_format.keep_with_next = True
    name = title.add_run("Ziyaaddin Yaramis")
    set_run_font(name, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(4)
    subtitle.paragraph_format.keep_with_next = True
    role = subtitle.add_run("Ohjelmistokehittäjä | ZyInnovan perustaja")
    set_run_font(role, size=12, color=ACCENT, bold=True)
    add_contact_line(document)

    add_section(document, "Ammatillinen tiivistelmä")
    summary = document.add_paragraph()
    summary_run = summary.add_run(
        "Olen ohjelmistokehittäjä ja ZyInnovan perustaja. Kehitän "
        "käytännönläheisiä verkko- ja mobiilituotteita, joissa painotan "
        "selkeää rajausta, ylläpidettävää toteutusta ja todennettuihin "
        "tietoihin perustuvaa viestintää. Projektikokemukseni kattaa "
        "Flutterin, Firebase Authenticationin, Cloud Firestoren, REST/HTTP-"
        "integraatiot sekä HTML-, CSS- ja JavaScript-kehityksen. Työskentelen "
        "englanniksi ja suomeksi B2-tasolla."
    )
    set_run_font(summary_run, size=10.2, color=BODY_COLOR)

    add_section(document, "Ydinosaaminen")
    for label, value in [
        ("Mobiilikehitys", "Flutter, Dart"),
        ("Web-kehitys", "HTML, CSS, JavaScript, Bootstrapin apuluokat"),
        ("Taustapalvelut ja data", "Firebase Authentication, Cloud Firestore"),
        ("Integraatiot", "REST/HTTP-rajapinnat, TMDB API"),
        ("Työskentelytavat", "Git, GitHub, dokumentointi, iteratiivinen kehitys"),
    ]:
        add_labeled_skill(document, label, value, bullet_num_id)

    add_section(document, "Kokemus")
    add_entry_title(
        document,
        "Perustaja ja ohjelmistokehittäjä - ZyInnova",
        "Nykyinen tehtävä",
    )
    for text in [
        "Määrittelen ja kehitän ZyInnovan alla perustajavetoisia verkko- ja "
        "mobiilituotteita tuoterajauksesta toteutukseen ja jatkokehitykseen.",
        "Ylläpidän julkista perustajaportfoliota ja pidän projektiviestinnän "
        "ajan tasalla. Tämänhetkisiin yksityisiin tuotehankkeisiin kuuluvat "
        "LemmikkiLife ja GameScan.",
    ]:
        add_bullet(document, text, bullet_num_id)
    add_project_link(
        document, "Verkkosivusto", "ZyInnova", "https://zyinnova.com/"
    )

    add_section(document, "Aiempi työkokemus")
    earlier_experience = [
        (
            "Flutter-kehittäjäharjoittelija - Roseance Oy",
            "loka-joulukuu 2024",
            "Suoritin dokumentoidun määräaikaisen harjoittelun "
            "Flutter-kehittäjäharjoittelijan tehtävässä.",
        ),
        (
            "Web-kehittäjäharjoittelija - Sisustusklinikka.fi",
            "helmi-toukokuu 2024",
            "Suoritin dokumentoidun määräaikaisen harjoittelun "
            "web-kehittäjäharjoittelijan tehtävässä.",
        ),
        (
            "Java-mentori - Wise Quarter",
            "2022",
            "Toimin Java-mentorina dokumentoidussa koulutusohjelmassa.",
        ),
        (
            "Poliisi - Turkin kansallinen poliisi",
            "2006-2016",
            "Työskentelin poliisina Turkin kansallisessa poliisissa.",
        ),
    ]
    for title_text, period, description in earlier_experience:
        add_entry_title(document, title_text, period)
        add_bullet(document, description, bullet_num_id)

    document.add_page_break()
    add_section(document, "Valitut projektit")

    add_entry_title(document, "LemmikkiLife", "Kehitteillä")
    add_bullet(
        document,
        "Kehitän yksityisessä kehitysvaiheessa olevaa ZyInnovan lemmikkien "
        "hoitoon tarkoitettua tuotetta, jonka tavoitteena on yhdistää hoidon "
        "seuranta, rokotus- ja lääkemuistutukset, tuotteet sekä tilauslaatikot "
        "yhteen sovellukseen.",
        bullet_num_id,
    )
    add_project_link(
        document,
        "Projektisivu",
        "LemmikkiLife-projektisivu",
        "https://zyinnova.com/projects/lemmikkilife",
    )

    add_entry_title(
        document, "Monikielinen Todo-sovellus", "Julkinen prototyyppi"
    )
    add_bullet(
        document,
        "Toteutin Flutter-sovelluksen, jossa on englannin- ja suomenkielinen "
        "lokalisointi sekä Firebase Authenticationiin perustuva kirjautuminen "
        "ja salasanan palautus. Käyttäjäkohtaiset tehtävät tallennetaan Cloud "
        "Firestoreen, ja ne tukevat prioriteetteja, määräpäiviä, suodatusta ja "
        "paikallisia ilmoituksia.",
        bullet_num_id,
    )
    add_project_link(
        document,
        "Lähdekoodi",
        "Todo-sovelluksen lähdekoodi",
        "https://github.com/ZiyaaddinYaramis/github_io_todo_app",
    )

    add_entry_title(document, "Perustajaportfolio", "Julkaistu")
    add_bullet(
        document,
        "Toteutin ja julkaisin henkilökohtaisen portfolioni HTML:llä, CSS:llä "
        "ja JavaScriptillä käyttäen Bootstrapin apuluokkia sekä AOS- ja "
        "Typed.js-kirjastoja. Ylläpidän sivustolla julkista "
        "ammattiprofiiliani.",
        bullet_num_id,
    )
    add_project_link(
        document,
        "Sivusto",
        "Portfolio",
        "https://portfolio.zyinnova.com/",
    )
    add_project_link(
        document,
        "Lähdekoodi",
        "Portfolion lähdekoodi",
        "https://github.com/ZiyaaddinYaramis/portfolio-zyinnova-portfolio",
    )

    add_section(document, "Koulutus")
    add_entry_title(
        document,
        "Tieto- ja viestintätekniikan perustutkinto",
        "Valmistunut 2025",
    )
    education = document.add_paragraph(style="CV Compact")
    institution = education.add_run("Keuda, Suomi")
    set_run_font(institution, size=10.0, color=BODY_COLOR, bold=True)
    qualification = education.add_run(
        " | Tutkintonimike: ohjelmistokehittäjä"
    )
    set_run_font(qualification, size=10.0, color=BODY_COLOR)

    add_section(document, "Lisäkoulutus")
    add_entry_title(
        document,
        "Full-Stack Java Developer -koulutus - TechPro Education",
        "Suoritettu 2021",
    )

    add_section(document, "Kielitaito")
    languages = document.add_paragraph(style="CV Compact")
    for index, (language, level) in enumerate(
        [("Turkki", "äidinkieli"), ("Englanti", "B2"), ("Suomi", "B2")]
    ):
        if index:
            separator = languages.add_run("  |  ")
            set_run_font(separator, size=10.2, color=MUTED)
        label = languages.add_run(f"{language}: ")
        set_run_font(label, size=10.2, color=NAVY, bold=True)
        value = languages.add_run(level)
        set_run_font(value, size=10.2, color=BODY_COLOR)

    return document


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = build_document()
    document.save(args.output)


if __name__ == "__main__":
    main()
